"""Generate comparative evaluation v2 Word document from markdown.

Layout matches MSALGCM-Final-Paper-Draft.docx:
  - Section 0 (single column): title, authors, emails, abstract, ACM metadata
  - Section 1+ (two columns): body from INTRODUCTION through appendices
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "Paper Setup" / "MSALGCM-Final-Paper-Draft.docx"
SOURCE_MD = ROOT / "Paper Setup" / "comparative evaluation v2.md"
OUTPUT = ROOT / "Paper Setup" / "comparative evaluation v2.docx"

AUTHORS = "Juan Miguel Miranda\tJulian Johan Briones\tLance Xavier Lim"
EMAILS = (
    "juan_miranda@dlsu.edu.ph\t"
    "Julian_briones@dlsu.edu.ph\t"
    "lance_xavier_lim@dlsu.edu.ph"
)
CATEGORIES = (
    "G.1.6 [Numerical Analysis]: Optimization — global optimization, "
    "stochastic programming; I.2.8 [Artificial Intelligence]: Problem "
    "Solving, Control Methods, and Search"
)
GENERAL_TERMS = "Algorithms, Performance, Experimentation, Measurement."


def style_map(doc: Document) -> dict[str, object]:
    """Map style display names to style objects (template uses custom style IDs)."""
    return {s.name: s for s in doc.styles}


def set_section_columns(section, num: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    if num == 1:
        if qn("w:num") in cols.attrib:
            del cols.attrib[qn("w:num")]
    else:
        cols.set(qn("w:num"), str(num))
        cols.set(qn("w:space"), "720")


def clear_body_keep_final_sectpr(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_para(doc: Document, text: str, style_name: str, styles: dict[str, object]) -> None:
    p = doc.add_paragraph()
    p.style = styles[style_name]
    if text:
        add_inline_formatting(p, text)


def add_blank(doc: Document, styles: dict[str, object]) -> None:
    add_para(doc, "", "Body Text", styles)


def add_text(doc: Document, text: str, styles: dict[str, object], style_name: str = "Body Text") -> None:
    add_para(doc, text, style_name, styles)


def add_inline_formatting(paragraph, text: str) -> None:
    """Add runs with **bold** and `code` markers."""
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_separator(line: str) -> bool:
    line = line.strip()
    if not line.startswith("|"):
        return False
    return bool(re.match(r"^\|[\s\-:|]+\|$", line))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.autofit = False
    usable_width = Inches(6.5)
    col_width = usable_width / n_cols
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell = table.rows[ri].cells[ci]
            cell.width = col_width
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_formatting(p, text)
            if ri == 0:
                for run in p.runs:
                    run.bold = True


def normalize_h1(text: str) -> str:
    """Strip leading 'N. ' from top-level section headings to match draft style."""
    m = re.match(r"^\d+\.\s+(.+)$", text.strip())
    return m.group(1) if m else text.strip()


def extract_front_matter(md_text: str) -> dict[str, str]:
    title_match = re.search(r"^#\s+(.+)$", md_text, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else ""

    abstract_match = re.search(
        r"## ABSTRACT\s*\n\n(.*?)\n\n---",
        md_text,
        re.DOTALL,
    )
    abstract_raw = abstract_match.group(1).strip() if abstract_match else ""
    abstract_paras: list[str] = []
    keywords = ""
    for block in re.split(r"\n\n+", abstract_raw):
        block = block.strip()
        if block.startswith("**Keywords:**"):
            keywords = block.replace("**Keywords:**", "").strip()
        else:
            abstract_paras.append(block)

    return {
        "title": title,
        "abstract_paras": abstract_paras,
        "keywords": keywords,
    }


def build_front_matter(doc: Document, meta: dict[str, str], styles: dict[str, object]) -> None:
    title_p = doc.add_paragraph()
    title_p.style = styles["Title"]
    title_p.add_run(meta["title"])
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_blank(doc, styles)
    add_para(doc, AUTHORS, "Heading 2", styles)
    add_blank(doc, styles)
    add_para(doc, EMAILS, "Normal", styles)
    add_blank(doc, styles)
    add_blank(doc, styles)
    add_para(doc, "ABSTRACT", "Heading 1", styles)
    for para in meta["abstract_paras"]:
        add_text(doc, para, styles, "Body Text")
    add_blank(doc, styles)
    add_para(doc, "Categories and Subject Descriptors", "Heading 2", styles)
    add_blank(doc, styles)
    add_para(doc, CATEGORIES, "Normal", styles)
    add_blank(doc, styles)
    add_para(doc, "General Terms", "Heading 2", styles)
    add_blank(doc, styles)
    add_para(doc, GENERAL_TERMS, "Body Text", styles)
    add_blank(doc, styles)
    add_para(doc, "Keywords", "Heading 2", styles)
    add_blank(doc, styles)
    add_para(doc, meta["keywords"], "Body Text", styles)
    add_blank(doc, styles)


def body_start_index(lines: list[str]) -> int:
    for i, line in enumerate(lines):
        if re.match(r"^##\s+1\.\s+INTRODUCTION\s*$", line.strip()):
            return i
    raise ValueError("Could not find '## 1. INTRODUCTION' in markdown source.")


def add_body_from_markdown(doc: Document, md_text: str, styles: dict[str, object]) -> None:
    lines = md_text.splitlines()
    start = body_start_index(lines)
    i = start
    table_buffer: list[list[str]] | None = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("|"):
            if table_buffer is None:
                table_buffer = []
            if not is_table_separator(stripped):
                table_buffer.append(parse_table_row(stripped))
            i += 1
            continue

        if table_buffer is not None:
            add_table(doc, table_buffer)
            table_buffer = None
            # do not increment i; reprocess current line

        if stripped.startswith("## "):
            heading = normalize_h1(stripped[3:])
            add_para(doc, heading, "Heading 1", styles)
            add_blank(doc, styles)
            i += 1
            continue

        if stripped.startswith("### "):
            add_para(doc, stripped[4:], "Heading 2", styles)
            i += 1
            continue

        if re.match(r"^(\d+\.\s|\*\s|-\s)", stripped):
            add_text(doc, stripped, styles, "List Paragraph")
            i += 1
            continue

        if stripped.startswith(">"):
            add_text(doc, stripped.lstrip("> ").strip(), styles, "Body Text")
            i += 1
            continue

        add_text(doc, stripped, styles, "Body Text")
        i += 1

    if table_buffer:
        add_table(doc, table_buffer)


def generate() -> Path:
    md_text = SOURCE_MD.read_text(encoding="utf-8")
    meta = extract_front_matter(md_text)

    shutil.copy(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))
    styles = style_map(doc)
    clear_body_keep_final_sectpr(doc)
    set_section_columns(doc.sections[0], 1)

    build_front_matter(doc, meta, styles)

    doc.add_section(WD_SECTION.CONTINUOUS)
    set_section_columns(doc.sections[-1], 2)

    add_body_from_markdown(doc, md_text, styles)

    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    out = generate()
    print(f"Wrote {out}")
